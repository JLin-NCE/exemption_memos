import docx
import pandas as pd
import os
from docx.shared import Pt
import re
import win32com.client

def format_location(text):
    """Format location text: extract between . and (, title case each word"""
    try:
        # Find text between . and (
        match = re.search(r'\.([^(]+)', text)
        if match:
            # Get the text and strip whitespace
            location = match.group(1).strip()
            # Title case each word
            return ' '.join(word.capitalize() for word in location.lower().split())
    except Exception as e:
        print(f"Error formatting location: {e}")
        return text
    return text

def process_checkboxes(doc_path):
    """Process checkboxes in the document using WIN32COM"""
    word_app = None
    doc = None
    
    # Define checkboxes to mark
    checkboxes_to_mark = {
        1: "Fire Hydrant",
        12: "Narrow Sidewalks",
        17: "Greater than 2.1% (1:48) transition area slope"
    }
    
    try:
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        doc = word_app.Documents.Open(os.path.abspath(doc_path))
        
        # Process specified checkboxes
        for checkbox_num, description in checkboxes_to_mark.items():
            field = doc.FormFields.Item(checkbox_num)
            if field.Type == 71:  # Verify it's a checkbox
                field.CheckBox.Value = 1
                print(f"Checked checkbox {checkbox_num}: {description}")
        
        doc.Save()
        print("Successfully processed checkboxes and saved document")
        
    except Exception as e:
        print(f"Error processing checkboxes: {e}")
    finally:
        try:
            if doc:
                doc.Close(False)
            if word_app:
                word_app.Quit()
        except Exception as e:
            print(f"Error closing Word application: {str(e)}")

def replace_first_instance(doc_path, excel_path, output_folder="output"):
    """Replace values while preserving labels"""
    print("\n=== Starting Document Processing ===")
    
    # Read first row from Excel (skipping header)
    print("Reading Excel values...")
    df = pd.read_excel(excel_path, skiprows=1)
    raw_location = str(df.iloc[0, 1])  # First row, column B
    location = format_location(raw_location)
    intersection = str(df.iloc[0, 2])  # First row, column C
    
    print(f"Values to insert:")
    print(f"Original location: '{raw_location}'")
    print(f"Formatted location: '{location}'")
    print(f"Intersection: '{intersection}'")
    
    # Create new document
    doc = docx.Document(doc_path)
    
    print("\nReplacing values and setting fonts...")
    # Replace first instances
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'CR location' in cell.text:
                    print(f"Found 'CR location' in cell: '{cell.text}'")
                    for para in cell.paragraphs:
                        if 'CR location' in para.text:
                            para.clear()
                            run = para.add_run(location)
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)
                            print(f"Replaced with: '{location}' in Times New Roman 12pt")
                
                if 'I location' in cell.text:
                    print(f"Found 'I location' in cell: '{cell.text}'")
                    for para in cell.paragraphs:
                        if 'I location' in para.text:
                            # Preserve "Intersection:" label
                            para.clear()
                            label_run = para.add_run("Intersection: ")
                            label_run.font.name = "Times New Roman"
                            label_run.font.size = Pt(12)
                            
                            value_run = para.add_run(intersection)
                            value_run.font.name = "Times New Roman"
                            value_run.font.size = Pt(12)
                            print(f"Replaced with: 'Intersection: {intersection}' in Times New Roman 12pt")
    
    # Save modified document
    os.makedirs(output_folder, exist_ok=True)
    output_path = os.path.join(output_folder, 'Location_001.docx')
    doc.save(output_path)
    print(f"\nSaved modified document to: {output_path}")
    
    # Process checkboxes after saving
    process_checkboxes(output_path)
    return output_path

if __name__ == "__main__":
    doc_path = "Curb Ramp Hardship Form - City of Redondo Beach.docx"
    excel_path = r"C:\Users\JLin\Downloads\Exemption Memo Code - Python\Redondo Beach - FY 22 Rehab - Curb Ramp Exemptions v02.xlsx"
    
    replace_first_instance(doc_path, excel_path)
